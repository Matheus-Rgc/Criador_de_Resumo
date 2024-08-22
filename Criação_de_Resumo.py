from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def Titulo_Terminal(a):
    print('=' * 30)
    print(a)
    print('=' * 30)
    return a

def linha():
    print('-' * 30)

def mostrar_topico(mostrar):
    if len(mostrar) == 0:
        return ' '
    else:
        return '\n'.join(mostrar)
    
def mostrar(mostrar):
    for c in mostrar:
        print(c)

def adicionar_paragrafo(doc, texto, estilo=None):
    paragrafo = doc.add_paragraph(texto, style=estilo)
    return paragrafo

# Criação do documento
doc = Document()

# Definição dos estilos
styles = doc.styles

# Estilo de parágrafo
p_style = styles.add_style('Paragraph', WD_STYLE_TYPE.PARAGRAPH)
p_style.font.name = 'Arial'
p_style.font.size = Pt(11)
p_style.font.bold = False

# Estilo do título principal
head_style = styles.add_style('Head', WD_STYLE_TYPE.PARAGRAPH)
head_style.font.name = 'Arial'
head_style.font.size = Pt(22)
head_style.font.color.rgb = RGBColor(0, 0, 0)
head_style.font.bold = True

# Estilo dos subtítulos
subhead_style = styles.add_style('SubHead', WD_STYLE_TYPE.PARAGRAPH)
subhead_style.font.name = 'Arial'
subhead_style.font.size = Pt(14)
subhead_style.font.color.rgb = RGBColor(0, 0, 255)

# Adicionando o título
Titulo_Terminal('     Criação de Resumo')
titulo = str(input('Sobre qual matéria é seu resumo: '))
linha()
doc.add_paragraph(titulo, style='Head')

# Coletando e adicionando os tópicos
pchaves = []
conteudo = []

a1 = int(input('Informe quantos tópicos gostaria de adicionar ao resumo: '))
linha()
for topic in range(a1):
    topico = input(f'Me diga seu {topic + 1}º tópico: ')
    pchaves.append(topico)
linha()

for p in pchaves:
    a2 = str(input(f'Defina e explique de maneira simples e rápida sobre {p}: '))
    conteudo.append(a2)
    linha()

# Adicionando os tópicos e suas explicações ao documento
for i, topico in enumerate(pchaves):
    doc.add_paragraph(topico, style='SubHead')
    doc.add_paragraph(conteudo[i], style='Paragraph')

# Salvando o documento com nome seguro
safe_title = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '', titulo)  # Remove caracteres inválidos
file_name = f'Resumo_sobre_{safe_title}.docx'
doc.save(file_name)

print(f"Documento '{file_name}' criado com sucesso!")

