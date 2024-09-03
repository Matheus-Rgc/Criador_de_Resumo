from docx import Document
from docx.shared import Inches
from datetime import datetime

def create_report():
    # Cria um novo documento Word
    doc = Document()

    # Adiciona informações no cabeçalho
    doc.add_paragraph('Campus: São Miguel'.ljust(50) + 'Turma: 6°J'.rjust(30))
    doc.add_paragraph('Disciplina: Estágio supervisionado específico em Processos Grupais')
    doc.add_paragraph('Estagiária: Sofia Mendonça dos Santos'.ljust(50) + 'RGM: 1650971-8'.rjust(30))
    doc.add_paragraph('Estagiária: Maria Fernanda Ribeiro'.ljust(50) + 'RGM: 267893-5'.rjust(30))
    doc.add_paragraph('Supervisora: Thais Baleeiro dos Santos'.ljust(50) + 'CRP: 06/55581'.rjust(30))
    doc.add_paragraph('Instituição:'.ljust(50))
    doc.add_paragraph('Data da Atividade: 07/03/2023'.ljust(50) + 'Relatório número: 01'.rjust(30))
    
    # Adiciona um espaço entre o cabeçalho e o corpo do texto
    doc.add_paragraph()

    # Adiciona a seção I - RELATO DA ATIVIDADE
    doc.add_paragraph('I. RELATO DA ATIVIDADE', style='Heading 1')
    
    doc.add_paragraph('1. Objetivo do encontro:')
    obj_encontro = input('Qual foi o objetivo do encontro? ')
    if not obj_encontro.strip():
        obj_encontro = 'Descrição não fornecida.'
    doc.add_paragraph(obj_encontro)
    
    doc.add_paragraph('2. Descrição:')
    descri = input('Descreva a visita: ')
    if not descri.strip():
        descri = 'Descrição não fornecida.'
    doc.add_paragraph(descri)

    # Adiciona a seção II - ANÁLISE CRÍTICA
    doc.add_paragraph('II. ANÁLISE CRÍTICA (fundamentação teórica facultativo do professor)', style='Heading 1')
    analise = input('Analise crítica: ')
    if not analise.strip():
        analise = 'Análise não fornecida.'
    doc.add_paragraph(analise)
    doc.add_paragraph(
        "Todos os primeiros contatos já conduzem a uma impressão preliminar de caráter diagnóstico, para o qual se deve conhecer também a história da instituição e - pelo menos - os grandes delineamentos de suas características. (BLEGER, 1984, p51)"
    )

    # Adiciona a seção III - IMPRESSÕES PESSOAIS
    doc.add_paragraph('III. IMPRESSÕES PESSOAIS (análise individual)', style='Heading 1')
    doc.add_paragraph('Matheus Gonçalves:')
    impre1 = input('Impressões pessoais de Matheus Gonçalves: ')
    if not impre1.strip():
        impre1 = 'Impressões não fornecidas.'
    doc.add_paragraph(impre1)
    doc.add_paragraph('Ana Faria:')
    impre2 = input('Impressões pessoais de Ana Faria: ')
    if not impre2.strip():
        impre2 = 'Impressões não fornecidas.'
    doc.add_paragraph(impre2)

    # Adiciona a seção IV - PROPOSTAS/ORIENTAÇÕES DE AÇÕES FUTURAS
    doc.add_paragraph('IV. PROPOSTAS/ORIENTAÇÕES DE AÇÕES FUTURAS (registro obrigatório do professor)', style='Heading 1')
    propostas = input('Propostas para a próxima visita: ')
    if not propostas.strip():
        propostas = 'Propostas não fornecidas.'
    doc.add_paragraph(propostas)

    # Adiciona a seção V - REFERÊNCIAS
    doc.add_paragraph('V. REFERÊNCIAS (se houver referencial teórico)', style='Heading 1')
    doc.add_paragraph(
        "BLEGER, J. Psico-higiene e psicologia institucional. Porto Alegre: Artes Médicas, 1984."
    )

    # Adiciona o local e a data
    doc.add_paragraph()
    doc.add_paragraph('São Paulo, 08 de Março de 2023.', style='Normal')

    # Salva o documento
    doc.save('Relatorio_Atividade.docx')

if __name__ == "__main__":
    create_report()

